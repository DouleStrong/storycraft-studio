from __future__ import annotations

import shutil
import subprocess
import tempfile
from html import escape
from pathlib import Path

from docx import Document
from docx.shared import Inches
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    HRFlowable,
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


CJK_FONT = "STSong-Light"
EMBEDDED_CJK_FONT = "StoryCraftWQY"
EMBEDDED_CJK_FONT_PATH = Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc")


def resolve_selected_chapters(project, selected_chapter_ids: list[int] | None = None) -> list:
    chapters = sorted(list(getattr(project, "chapters", []) or []), key=lambda item: getattr(item, "order_index", 0))
    if not selected_chapter_ids:
        return chapters

    selected_ids = set(selected_chapter_ids)
    filtered = [chapter for chapter in chapters if getattr(chapter, "id", None) in selected_ids]
    return filtered or chapters


def resolve_selected_illustrations(scene, selected_illustration_ids: list[int] | None = None) -> list:
    illustrations = sorted(
        list(getattr(scene, "illustrations", []) or []),
        key=lambda item: (
            0 if getattr(item, "is_canonical", False) else 1,
            getattr(item, "candidate_index", 0),
            getattr(item, "id", 0) or 0,
        ),
    )
    if not illustrations:
        return []

    if selected_illustration_ids:
        selected_ids = set(selected_illustration_ids)
        filtered = [item for item in illustrations if getattr(item, "id", None) in selected_ids]
        if filtered:
            return filtered[:1]

    canonical = [item for item in illustrations if getattr(item, "is_canonical", False)]
    return (canonical or illustrations)[:1]


def build_export_delivery_summary(project, export_package, serialized_files: list[dict]) -> dict:
    chapters = resolve_selected_chapters(project, getattr(export_package, "selected_chapter_ids", None))
    selected_illustration_ids = getattr(export_package, "selected_illustration_ids", None) or []
    if selected_illustration_ids:
        illustration_count = len(selected_illustration_ids)
    else:
        illustration_count = 0
        for chapter in chapters:
            for scene in getattr(chapter, "scenes", []) or []:
                illustration_count += len(resolve_selected_illustrations(scene, None))

    quality_states = [
        str(file_info.get("quality_check", {}).get("status", "")).strip()
        for file_info in serialized_files
        if file_info.get("quality_check")
    ]
    if "failed" in quality_states:
        quality_status = "failed"
    elif "warn" in quality_states:
        quality_status = "warn"
    elif "passed" in quality_states:
        quality_status = "passed"
    else:
        quality_status = "pending"

    total_size_bytes = sum(int(file_info.get("size_bytes") or 0) for file_info in serialized_files)
    total_page_count = sum(int(file_info.get("page_count") or 0) for file_info in serialized_files)

    return {
        "project_title": getattr(project, "title", ""),
        "chapter_count": len(chapters),
        "character_count": len(list(getattr(project, "characters", []) or [])),
        "illustration_count": illustration_count,
        "quality_status": quality_status,
        "total_size_bytes": total_size_bytes,
        "total_page_count": total_page_count,
        "ready_formats": [str(file_info.get("format", "")).upper() for file_info in serialized_files if file_info.get("path")],
    }


def build_pdf_bundle(
    project,
    target_path: Path,
    *,
    selected_chapter_ids: list[int] | None = None,
    selected_illustration_ids: list[int] | None = None,
) -> dict:
    _ensure_cjk_font()
    chapters = resolve_selected_chapters(project, selected_chapter_ids)
    styles = _build_pdf_styles()
    project_title = str(getattr(project, "title", "未命名作品")).strip() or "未命名作品"
    doc = SimpleDocTemplate(
        str(target_path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=16 * mm,
        title=project_title,
    )

    story = []
    story.extend(_build_cover_elements(project, styles))
    story.extend(_build_character_elements(project, styles))
    story.extend(_build_chapter_elements(chapters, styles, selected_illustration_ids))

    doc.build(
        story,
        onFirstPage=lambda canvas, document: _draw_cover_page_chrome(canvas, document, project_title=project_title),
        onLaterPages=lambda canvas, document: _draw_reading_page_chrome(canvas, document, project_title=project_title),
    )
    quality_check = validate_pdf_bundle(target_path, project, chapters)
    if quality_check["status"] == "failed":
        raise ValueError("Generated PDF failed quality checks")

    return {
        "path": str(target_path),
        "filename": target_path.name,
        "size_bytes": target_path.stat().st_size,
        "page_count": quality_check["page_count"],
        "quality_check": quality_check,
    }


def build_docx_bundle(
    project,
    target_path: Path,
    *,
    selected_chapter_ids: list[int] | None = None,
    selected_illustration_ids: list[int] | None = None,
) -> dict:
    chapters = resolve_selected_chapters(project, selected_chapter_ids)
    document = Document()
    document.add_heading(getattr(project, "title", "未命名作品"), level=0)
    document.add_paragraph(f"{getattr(project, 'genre', '')} | {getattr(project, 'era', '')}")
    document.add_paragraph(getattr(project, "logline", ""))
    document.add_paragraph(f"基调：{getattr(project, 'tone', '')}")
    document.add_paragraph(f"目标篇幅：{getattr(project, 'target_length', '')}")

    document.add_page_break()
    document.add_heading("角色档案", level=1)
    for character in list(getattr(project, "characters", []) or []):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{getattr(character, 'name', '')}").bold = True
        paragraph.add_run(f" / {getattr(character, 'role', '')}")
        document.add_paragraph(getattr(character, "signature_line", "") or getattr(character, "personality", ""))
        document.add_paragraph(f"目标：{getattr(character, 'goal', '')}")
        document.add_paragraph(f"外貌：{getattr(character, 'appearance', '')}")
        document.add_paragraph(f"口吻：{getattr(character, 'speech_style', '')}")

    for chapter in chapters:
        document.add_page_break()
        document.add_heading(getattr(chapter, "title", "未命名章节"), level=1)
        document.add_paragraph(getattr(chapter, "summary", ""))
        document.add_paragraph(f"章节目标：{getattr(chapter, 'chapter_goal', '')}")
        document.add_paragraph(f"章节钩子：{getattr(chapter, 'hook', '')}")
        for block in sorted(list(getattr(chapter, "narrative_blocks", []) or []), key=lambda item: getattr(item, "order_index", 0)):
            document.add_paragraph(getattr(block, "content", ""))
        for scene in sorted(list(getattr(chapter, "scenes", []) or []), key=lambda item: getattr(item, "order_index", 0)):
            document.add_heading(
                f"{getattr(scene, 'scene_type', '')} {getattr(scene, 'location', '')} - {getattr(scene, 'time_of_day', '')}",
                level=2,
            )
            document.add_paragraph(f"场景目标：{getattr(scene, 'objective', '')}")
            document.add_paragraph(f"情绪：{getattr(scene, 'emotional_tone', '')}")
            for dialogue in sorted(list(getattr(scene, "dialogue_blocks", []) or []), key=lambda item: getattr(item, "order_index", 0)):
                parenthetical = getattr(dialogue, "parenthetical", "")
                if parenthetical:
                    line = f"{getattr(dialogue, 'speaker', '')}（{parenthetical}）：{getattr(dialogue, 'content', '')}"
                else:
                    line = f"{getattr(dialogue, 'speaker', '')}：{getattr(dialogue, 'content', '')}"
                document.add_paragraph(line)

            for illustration in resolve_selected_illustrations(scene, selected_illustration_ids):
                image_path = Path(str(getattr(illustration, "file_path", "")))
                if image_path.exists():
                    document.add_picture(str(image_path), width=Inches(5.5))
                    document.add_paragraph("关键剧照")

    document.save(str(target_path))
    return {
        "path": str(target_path),
        "filename": target_path.name,
        "size_bytes": target_path.stat().st_size,
    }


def validate_pdf_bundle(target_path: Path, project, chapters: list) -> dict:
    checks = []

    def add_check(*, key: str, passed: bool, passed_message: str, failed_message: str, severity: str = "critical") -> None:
        checks.append(
            {
                "key": key,
                "status": "passed" if passed else ("failed" if severity == "critical" else "warn"),
                "message": passed_message if passed else failed_message,
                "severity": severity,
            }
        )

    add_check(
        key="file_exists",
        passed=target_path.exists(),
        passed_message="PDF 文件已生成。",
        failed_message="PDF 文件不存在。",
    )
    if not target_path.exists():
        return {"status": "failed", "checks": checks, "page_count": 0, "render_check": {"status": "skipped", "message": "PDF 不存在。"}}

    file_size = target_path.stat().st_size
    add_check(
        key="non_empty",
        passed=file_size > 0,
        passed_message="PDF 文件大小正常。",
        failed_message="PDF 文件为空。",
    )

    reader = PdfReader(str(target_path))
    extracted_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    page_count = len(reader.pages)
    add_check(
        key="has_pages",
        passed=page_count > 0,
        passed_message=f"PDF 共 {page_count} 页。",
        failed_message="PDF 没有可读取页。",
    )
    project_title = str(getattr(project, "title", "")).strip()
    add_check(
        key="title_present",
        passed=bool(project_title and project_title in extracted_text),
        passed_message="已检测到作品标题。",
        failed_message="未在 PDF 中检测到作品标题。",
    )
    add_check(
        key="character_section_present",
        passed="角色档案" in extracted_text,
        passed_message="已检测到角色页。",
        failed_message="未在 PDF 中检测到角色页。",
        severity="warning",
    )
    if chapters:
        chapter_title = str(getattr(chapters[0], "title", "")).strip()
        add_check(
            key="chapter_heading_present",
            passed=bool(chapter_title and chapter_title in extracted_text),
            passed_message=f"已检测到章节标题“{chapter_title}”。",
            failed_message="未在 PDF 中检测到章节标题。",
        )

    render_check = _render_pdf_preview_check(target_path)
    statuses = [check["status"] for check in checks]
    if "failed" in statuses:
        status = "failed"
    elif "warn" in statuses:
        status = "warn"
    else:
        status = "passed"
    return {
        "status": status,
        "checks": checks,
        "page_count": page_count,
        "render_check": render_check,
    }


def _build_cover_elements(project, styles) -> list:
    font_name = _cjk_font_name()
    summary_table = Table(
        [
            ["题材", getattr(project, "genre", "")],
            ["时代", getattr(project, "era", "")],
            ["气质", getattr(project, "tone", "")],
            ["目标篇幅", getattr(project, "target_length", "")],
        ],
        colWidths=[24 * mm, 128 * mm],
    )
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.Color(0.975, 0.975, 0.975, alpha=1)),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 11),
                ("GRID", (0, 0), (-1, -1), 0.7, colors.Color(0, 0, 0, alpha=0.18)),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
            ]
        )
    )

    return [
        Spacer(1, 42 * mm),
        Paragraph("StoryCraft Studio", styles["CoverEyebrow"]),
        Spacer(1, 10 * mm),
        Paragraph(_safe_html(getattr(project, "title", "未命名作品")), styles["CoverTitle"]),
        Spacer(1, 8 * mm),
        Paragraph(_safe_html(getattr(project, "logline", "")), styles["CoverLead"]),
        Spacer(1, 18 * mm),
        summary_table,
        Spacer(1, 16 * mm),
        Paragraph("图文审查版", styles["CoverTag"]),
        PageBreak(),
    ]


def _build_character_elements(project, styles) -> list:
    story = [Paragraph("角色档案", styles["SectionTitle"]), Spacer(1, 6 * mm)]
    for character in list(getattr(project, "characters", []) or []):
        card = Table(
            [
                [Paragraph(_safe_html(f"{getattr(character, 'name', '')} · {getattr(character, 'role', '')}"), styles["CharacterName"])],
                [Paragraph(_safe_html(getattr(character, "signature_line", "") or getattr(character, "personality", "")), styles["CharacterCardBody"])],
                [Paragraph(_safe_html(f"目标：{getattr(character, 'goal', '')}"), styles["CharacterCardMeta"])],
                [Paragraph(_safe_html(f"外貌：{getattr(character, 'appearance', '')}"), styles["CharacterCardMeta"])],
                [Paragraph(_safe_html(f"口吻：{getattr(character, 'speech_style', '')}"), styles["CharacterCardMeta"])],
            ],
            colWidths=[170 * mm],
        )
        card.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.Color(0.97, 0.97, 0.97, alpha=1)),
                    ("BOX", (0, 0), (-1, -1), 0.9, colors.Color(0, 0, 0, alpha=0.18)),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING", (0, 0), (-1, -1), 12),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ]
            )
        )
        story.extend([card, Spacer(1, 5 * mm)])
    if len(story) == 2:
        story.append(Paragraph("当前作品暂无角色。", styles["PageMeta"]))
    story.append(PageBreak())
    return story


def _build_chapter_elements(chapters, styles, selected_illustration_ids: list[int] | None) -> list:
    story = []
    for chapter_index, chapter in enumerate(chapters):
        if chapter_index:
            story.append(PageBreak())
        story.extend(
            [
                Paragraph(f"Chapter {chapter_index + 1:02d}", styles["ChapterEyebrow"]),
                Spacer(1, 2.2 * mm),
                HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#b68e4a"), lineCap="round"),
                Spacer(1, 5 * mm),
                Paragraph(_safe_html(getattr(chapter, "title", "未命名章节")), styles["ChapterTitle"]),
                Spacer(1, 4 * mm),
                Paragraph(_safe_html(getattr(chapter, "summary", "")), styles["Lead"]),
                Spacer(1, 4 * mm),
                Paragraph(_safe_html(f"章节目标：{getattr(chapter, 'chapter_goal', '')}"), styles["PageMeta"]),
                Paragraph(_safe_html(f"章节钩子：{getattr(chapter, 'hook', '')}"), styles["PageMeta"]),
                Spacer(1, 6 * mm),
            ]
        )

        for block in sorted(list(getattr(chapter, "narrative_blocks", []) or []), key=lambda item: getattr(item, "order_index", 0)):
            story.extend([Paragraph(_safe_html(getattr(block, "content", "")), styles["Body"]), Spacer(1, 3 * mm)])

        scenes = sorted(list(getattr(chapter, "scenes", []) or []), key=lambda item: getattr(item, "order_index", 0))
        if scenes:
            story.extend(
                [
                    Spacer(1, 4 * mm),
                    HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbb58e")),
                    Spacer(1, 3 * mm),
                    Paragraph("场景与关键画面", styles["SectionTitleMinor"]),
                    Spacer(1, 3 * mm),
                ]
            )
        for scene in scenes:
            story.append(
                Paragraph(
                    _safe_html(
                        f"{getattr(scene, 'scene_type', '')} {getattr(scene, 'location', '')} / {getattr(scene, 'time_of_day', '')}"
                    ),
                    styles["SceneMeta"],
                )
            )
            story.append(Paragraph(_safe_html(f"场景目标：{getattr(scene, 'objective', '')}"), styles["PageMeta"]))
            story.append(Paragraph(_safe_html(f"情绪：{getattr(scene, 'emotional_tone', '')}"), styles["PageMeta"]))
            story.append(Spacer(1, 2 * mm))

            for dialogue in sorted(list(getattr(scene, "dialogue_blocks", []) or []), key=lambda item: getattr(item, "order_index", 0)):
                parenthetical = f"（{getattr(dialogue, 'parenthetical', '')}）" if getattr(dialogue, "parenthetical", "") else ""
                dialogue_text = f"<b>{escape(getattr(dialogue, 'speaker', ''))}</b>{escape(parenthetical)}：{escape(getattr(dialogue, 'content', ''))}"
                story.append(Paragraph(dialogue_text, styles["Dialogue"]))

            for illustration in resolve_selected_illustrations(scene, selected_illustration_ids):
                image_path = Path(str(getattr(illustration, "file_path", "")))
                if image_path.exists():
                    story.extend(
                        [
                            Spacer(1, 3 * mm),
                            _build_scaled_pdf_image(image_path),
                            Spacer(1, 2 * mm),
                            Paragraph("关键剧照", styles["ImageCaption"]),
                        ]
                    )
            story.append(Spacer(1, 5 * mm))
    return story or [Paragraph("当前作品还没有可导出的章节内容。", styles["Body"])]


def _build_scaled_pdf_image(image_path: Path):
    image = RLImage(str(image_path))
    image.hAlign = "CENTER"
    image._restrictSize(164 * mm, 92 * mm)
    return image


def _build_pdf_styles():
    font_name = _cjk_font_name()
    base = getSampleStyleSheet()
    ink = colors.black
    styles = {
        "CoverEyebrow": ParagraphStyle(
            "CoverEyebrow",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=11,
            textColor=ink,
            alignment=TA_CENTER,
            leading=14,
        ),
        "CoverTitle": ParagraphStyle(
            "CoverTitle",
            parent=base["Title"],
            fontName=font_name,
            fontSize=28,
            leading=36,
            alignment=TA_CENTER,
            textColor=ink,
        ),
        "CoverLead": ParagraphStyle(
            "CoverLead",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=13,
            leading=20,
            alignment=TA_CENTER,
            textColor=ink,
        ),
        "CoverTag": ParagraphStyle(
            "CoverTag",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=11,
            leading=14,
            alignment=TA_CENTER,
            textColor=ink,
        ),
        "SectionTitle": ParagraphStyle(
            "SectionTitle",
            parent=base["Heading1"],
            fontName=font_name,
            fontSize=20,
            leading=25,
            textColor=ink,
        ),
        "SectionTitleMinor": ParagraphStyle(
            "SectionTitleMinor",
            parent=base["Heading2"],
            fontName=font_name,
            fontSize=16,
            leading=20,
            textColor=ink,
        ),
        "ChapterEyebrow": ParagraphStyle(
            "ChapterEyebrow",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=9.8,
            leading=12,
            textColor=ink,
        ),
        "ChapterTitle": ParagraphStyle(
            "ChapterTitle",
            parent=base["Heading1"],
            fontName=font_name,
            fontSize=24,
            leading=31,
            textColor=ink,
        ),
        "Lead": ParagraphStyle(
            "Lead",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=12.5,
            leading=20,
            textColor=ink,
        ),
        "Body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=11.5,
            leading=19,
            textColor=ink,
            spaceAfter=2 * mm,
        ),
        "SceneMeta": ParagraphStyle(
            "SceneMeta",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=12,
            leading=16,
            textColor=ink,
        ),
        "Dialogue": ParagraphStyle(
            "Dialogue",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=11.2,
            leading=18,
            leftIndent=4 * mm,
            textColor=ink,
        ),
        "PageMeta": ParagraphStyle(
            "PageMeta",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=10.8,
            leading=16,
            textColor=ink,
        ),
        "CharacterCardBody": ParagraphStyle(
            "CharacterCardBody",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=11.2,
            leading=17,
            textColor=ink,
        ),
        "CharacterName": ParagraphStyle(
            "CharacterName",
            parent=base["Heading2"],
            fontName=font_name,
            fontSize=15.5,
            leading=20,
            textColor=ink,
        ),
        "CharacterCardMeta": ParagraphStyle(
            "CharacterCardMeta",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=10.8,
            leading=16,
            textColor=ink,
        ),
        "ImageCaption": ParagraphStyle(
            "ImageCaption",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=9.8,
            leading=13,
            alignment=TA_CENTER,
            textColor=ink,
        ),
    }
    # Backward-compatible aliases for older internal callers.
    styles["CardBody"] = styles["CharacterCardBody"]
    styles["CardMeta"] = styles["CharacterCardMeta"]
    return styles


def _draw_cover_page_chrome(pdf_canvas, document, *, project_title: str) -> None:
    font_name = _cjk_font_name()
    pdf_canvas.saveState()
    pdf_canvas.setStrokeColor(colors.Color(0, 0, 0, alpha=0.18))
    pdf_canvas.setLineWidth(1)
    inset = 13 * mm
    pdf_canvas.roundRect(
        inset,
        inset,
        document.pagesize[0] - (2 * inset),
        document.pagesize[1] - (2 * inset),
        7 * mm,
        stroke=1,
        fill=0,
    )
    pdf_canvas.setFont(font_name, 8.8)
    pdf_canvas.setFillColor(colors.black)
    pdf_canvas.drawString(document.leftMargin, document.pagesize[1] - 12 * mm, "StoryCraft Studio")
    pdf_canvas.drawRightString(document.pagesize[0] - document.rightMargin, document.pagesize[1] - 12 * mm, project_title)
    pdf_canvas.restoreState()


def _draw_reading_page_chrome(pdf_canvas, document, *, project_title: str) -> None:
    font_name = _cjk_font_name()
    page_number = pdf_canvas.getPageNumber()
    pdf_canvas.saveState()
    header_y = document.pagesize[1] - 11 * mm
    rule_y = document.pagesize[1] - 15 * mm
    footer_rule_y = 13.5 * mm
    footer_text_y = 8.5 * mm

    pdf_canvas.setFont(font_name, 8.8)
    pdf_canvas.setFillColor(colors.black)
    pdf_canvas.drawString(document.leftMargin, header_y, project_title)
    pdf_canvas.drawRightString(document.pagesize[0] - document.rightMargin, header_y, "图文审查版")

    pdf_canvas.setStrokeColor(colors.Color(0, 0, 0, alpha=0.16))
    pdf_canvas.setLineWidth(0.6)
    pdf_canvas.line(document.leftMargin, rule_y, document.pagesize[0] - document.rightMargin, rule_y)
    pdf_canvas.line(document.leftMargin, footer_rule_y, document.pagesize[0] - document.rightMargin, footer_rule_y)

    pdf_canvas.setFont(font_name, 9)
    pdf_canvas.setFillColor(colors.black)
    pdf_canvas.drawRightString(document.pagesize[0] - document.rightMargin, footer_text_y, str(page_number))
    pdf_canvas.restoreState()


def _safe_html(value) -> str:
    return escape(str(value or "")).replace("\n", "<br/>")


def _cjk_font_name() -> str:
    return _ensure_cjk_font()


def _ensure_cjk_font() -> str:
    if EMBEDDED_CJK_FONT_PATH.exists():
        try:
            pdfmetrics.getFont(EMBEDDED_CJK_FONT)
        except KeyError:
            pdfmetrics.registerFont(TTFont(EMBEDDED_CJK_FONT, str(EMBEDDED_CJK_FONT_PATH), subfontIndex=0))
        return EMBEDDED_CJK_FONT

    try:
        pdfmetrics.getFont(CJK_FONT)
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont(CJK_FONT))
    return CJK_FONT


def _render_pdf_preview_check(target_path: Path) -> dict:
    if shutil.which("pdftoppm") is None:
        return {"status": "skipped", "message": "当前环境未安装 pdftoppm，已跳过渲染级校验。"}

    with tempfile.TemporaryDirectory(prefix="storycraft-pdf-check-") as tmpdir:
        output_prefix = Path(tmpdir) / "page"
        completed = subprocess.run(
            ["pdftoppm", "-f", "1", "-l", "3", "-png", str(target_path), str(output_prefix)],
            capture_output=True,
            text=True,
            check=False,
        )
        preview_paths = sorted(Path(tmpdir).glob("page-*.png"))
        if completed.returncode == 0 and preview_paths and all(path.stat().st_size > 0 for path in preview_paths):
            return {"status": "passed", "message": f"前 {len(preview_paths)} 页渲染检查通过。"}
        stderr = completed.stderr.strip() or "未知错误"
        return {"status": "warn", "message": f"渲染级校验未通过：{stderr}"}
